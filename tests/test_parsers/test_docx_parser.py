"""Tests for DOCX document parser."""

import pytest
from unittest.mock import Mock, patch, MagicMock
import tempfile
from pathlib import Path
import io

from src.parsers.docx_parser import DOCXParser
from src.parsers.base import ParsingError, ReconstructionError
from src.models.document import (
    DocumentStructure, PageStructure, TextRegion, VisualElement,
    BoundingBox, Dimensions, TextFormatting, DocumentMetadata
)


class TestDOCXParser:
    """Test cases for DOCXParser."""
    
    def test_initialization(self):
        """Test DOCX parser initialization."""
        parser = DOCXParser()
        
        assert parser.get_supported_formats() == ["docx"]
        assert hasattr(parser, 'logger')
        assert parser.logger.name == 'DOCXParser'
        assert parser.default_page_width == 612
        assert parser.default_page_height == 792
    
    def test_get_supported_formats(self):
        """Test getting supported formats."""
        parser = DOCXParser()
        formats = parser.get_supported_formats()
        
        assert formats == ["docx"]
        assert isinstance(formats, list)
    
    @patch('docx.Document')
    def test_parse_success(self, mock_document_class):
        """Test successful DOCX parsing."""
        parser = DOCXParser()
        
        # Mock Document object
        mock_doc = Mock()
        mock_document_class.return_value = mock_doc
        
        # Mock core properties
        mock_core_props = Mock()
        mock_core_props.title = "Test Document"
        mock_core_props.author = "Test Author"
        mock_core_props.subject = "Test Subject"
        mock_core_props.creator = "Test Creator"
        mock_core_props.created = None
        mock_core_props.modified = None
        mock_doc.core_properties = mock_core_props
        
        # Mock paragraphs
        mock_paragraph = Mock()
        mock_paragraph.text = "Test paragraph text"
        
        mock_run = Mock()
        mock_run.text = "Test paragraph text"
        mock_run.font = Mock()
        mock_run.font.size = None
        mock_run.font.name = "Calibri"
        mock_run.font.bold = False
        mock_run.font.italic = False
        mock_run.font.underline = False
        mock_run.font.color = None
        
        mock_paragraph.runs = [mock_run]
        mock_paragraph.alignment = None
        mock_doc.paragraphs = [mock_paragraph]
        
        # Mock tables and relationships
        mock_doc.tables = []
        mock_doc.part.rels.values.return_value = []
        
        # Mock file stat
        with patch('pathlib.Path.stat') as mock_stat:
            mock_stat.return_value.st_size = 2048
            
            structure = parser.parse("test.docx")
        
        assert isinstance(structure, DocumentStructure)
        assert structure.format == "docx"
        assert len(structure.pages) == 1
        assert structure.metadata.title == "Test Document"
        assert structure.metadata.author == "Test Author"
        assert structure.metadata.file_size == 2048
        
        # Check that text regions were created
        page = structure.pages[0]
        assert len(page.text_regions) == 1
        assert page.text_regions[0].text_content == "Test paragraph text"
    
    @patch('docx.Document')
    def test_parse_failure(self, mock_document_class):
        """Test DOCX parsing failure."""
        parser = DOCXParser()
        
        # Mock exception during document opening
        mock_document_class.side_effect = Exception("Failed to open document")
        
        with pytest.raises(ParsingError, match="Failed to parse DOCX document"):
            parser.parse("invalid.docx")
    
    def test_extract_run_formatting(self):
        """Test run formatting extraction."""
        parser = DOCXParser()
        
        # Mock run with formatting
        mock_run = Mock()
        mock_font = Mock()
        mock_font.size = Mock()
        mock_font.size.pt = 14
        mock_font.name = "Times New Roman"
        mock_font.bold = True
        mock_font.italic = True
        mock_font.underline = True
        
        # Mock color
        mock_color = Mock()
        mock_rgb = Mock()
        mock_rgb.r = 255
        mock_rgb.g = 0
        mock_rgb.b = 0
        mock_color.rgb = mock_rgb
        mock_font.color = mock_color
        
        mock_run.font = mock_font
        
        # Mock paragraph alignment
        mock_paragraph = Mock()
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        mock_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        formatting = parser._extract_run_formatting(mock_run, mock_paragraph)
        
        assert formatting.font_family == "Times New Roman"
        assert formatting.font_size == 14
        assert formatting.is_bold is True
        assert formatting.is_italic is True
        assert formatting.is_underlined is True
        assert formatting.color == "#FF0000"
        assert formatting.alignment == "center"
    
    def test_extract_run_formatting_defaults(self):
        """Test run formatting extraction with defaults."""
        parser = DOCXParser()
        
        # Mock run with minimal formatting
        mock_run = Mock()
        mock_font = Mock()
        mock_font.size = None
        mock_font.name = None
        mock_font.bold = None
        mock_font.italic = None
        mock_font.underline = None
        mock_font.color = None
        mock_run.font = mock_font
        
        # Mock paragraph with default alignment
        mock_paragraph = Mock()
        mock_paragraph.alignment = None
        
        formatting = parser._extract_run_formatting(mock_run, mock_paragraph)
        
        assert formatting.font_family == "Calibri"
        assert formatting.font_size == 12
        assert formatting.is_bold is False
        assert formatting.is_italic is False
        assert formatting.is_underlined is False
        assert formatting.color == "#000000"
        assert formatting.alignment == "left"
    
    def test_process_paragraph(self):
        """Test paragraph processing."""
        parser = DOCXParser()
        
        # Mock paragraph with runs
        mock_paragraph = Mock()
        
        # Create mock runs
        mock_run1 = Mock()
        mock_run1.text = "Hello "
        mock_run1.font = Mock()
        mock_run1.font.size = None
        mock_run1.font.name = "Arial"
        mock_run1.font.bold = False
        mock_run1.font.italic = False
        mock_run1.font.underline = False
        mock_run1.font.color = None
        
        mock_run2 = Mock()
        mock_run2.text = "World"
        mock_run2.font = Mock()
        mock_run2.font.size = None
        mock_run2.font.name = "Arial"
        mock_run2.font.bold = True
        mock_run2.font.italic = False
        mock_run2.font.underline = False
        mock_run2.font.color = None
        
        mock_paragraph.runs = [mock_run1, mock_run2]
        mock_paragraph.alignment = None
        
        text_regions = parser._process_paragraph(mock_paragraph, 0, 100)
        
        assert len(text_regions) == 2
        assert text_regions[0].text_content == "Hello "
        assert text_regions[1].text_content == "World"
        assert text_regions[0].formatting.is_bold is False
        assert text_regions[1].formatting.is_bold is True
        assert text_regions[0].bounding_box.y == 100
        assert text_regions[1].bounding_box.y == 100
    
    def test_process_table(self):
        """Test table processing."""
        parser = DOCXParser()
        
        # Mock table with cells
        mock_table = Mock()
        
        # Mock row with cells
        mock_cell1 = Mock()
        mock_cell1.text = "Cell 1"
        mock_cell2 = Mock()
        mock_cell2.text = "Cell 2"
        
        mock_row = Mock()
        mock_row.cells = [mock_cell1, mock_cell2]
        
        mock_table.rows = [mock_row]
        
        text_regions = parser._process_table(mock_table, 0, 200)
        
        assert len(text_regions) == 2
        assert text_regions[0].text_content == "Cell 1"
        assert text_regions[1].text_content == "Cell 2"
        assert text_regions[0].bounding_box.x == 50  # First column
        assert text_regions[1].bounding_box.x == 150  # Second column
        assert text_regions[0].bounding_box.y == 200
        assert text_regions[1].bounding_box.y == 200
    
    def test_calculate_distance(self):
        """Test distance calculation between bounding boxes."""
        parser = DOCXParser()
        
        bbox1 = BoundingBox(x=0, y=0, width=10, height=10)
        bbox2 = BoundingBox(x=10, y=0, width=10, height=10)
        
        # Distance between centers: (5,5) and (15,5) = 10
        distance = parser._calculate_distance(bbox1, bbox2)
        assert distance == 10.0
        
        # Same box should have distance 0
        distance = parser._calculate_distance(bbox1, bbox1)
        assert distance == 0.0
    
    def test_group_text_regions_by_position(self):
        """Test grouping text regions by position."""
        parser = DOCXParser()
        
        # Create text regions at different positions
        region1 = TextRegion(
            id="1",
            bounding_box=BoundingBox(x=0, y=10, width=50, height=20),
            text_content="Line 1 part 1"
        )
        region2 = TextRegion(
            id="2",
            bounding_box=BoundingBox(x=50, y=15, width=50, height=20),
            text_content="Line 1 part 2"
        )
        region3 = TextRegion(
            id="3",
            bounding_box=BoundingBox(x=0, y=50, width=100, height=20),
            text_content="Line 2"
        )
        
        groups = parser._group_text_regions_by_position([region1, region2, region3])
        
        assert len(groups) == 2
        assert len(groups[0]) == 2  # First line has 2 parts
        assert len(groups[1]) == 1  # Second line has 1 part
        assert groups[0][0].text_content == "Line 1 part 1"
        assert groups[0][1].text_content == "Line 1 part 2"
        assert groups[1][0].text_content == "Line 2"
    
    @patch('docx.Document')
    def test_reconstruct_success(self, mock_document_class):
        """Test successful DOCX reconstruction."""
        parser = DOCXParser()
        
        # Mock new document
        mock_doc = Mock()
        mock_document_class.return_value = mock_doc
        
        # Mock core properties
        mock_core_props = Mock()
        mock_doc.core_properties = mock_core_props
        
        # Mock paragraph and run creation
        mock_paragraph = Mock()
        mock_run = Mock()
        mock_run.font = Mock()
        mock_paragraph.add_run.return_value = mock_run
        mock_doc.add_paragraph.return_value = mock_paragraph
        
        # Create test document structure
        structure = DocumentStructure(format="docx")
        metadata = DocumentMetadata(title="Test Doc", author="Test Author")
        structure.metadata = metadata
        
        dims = Dimensions(width=612, height=792)
        page = PageStructure(page_number=1, dimensions=dims)
        
        # Add text region
        text_region = TextRegion(
            bounding_box=BoundingBox(x=50, y=50, width=200, height=20),
            text_content="Test text",
            formatting=TextFormatting(font_size=12.0, color="#000000")
        )
        page.text_regions.append(text_region)
        
        structure.add_page(page)
        
        # Mock document_to_bytes
        with patch.object(parser, '_document_to_bytes', return_value=b"docx content"):
            result = parser.reconstruct(structure)
        
        assert result == b"docx content"
        mock_doc.add_paragraph.assert_called()
        mock_paragraph.add_run.assert_called_with("Test text")
    
    @patch('docx.Document')
    def test_reconstruct_failure(self, mock_document_class):
        """Test DOCX reconstruction failure."""
        parser = DOCXParser()
        
        # Mock exception during document creation
        mock_document_class.side_effect = Exception("Reconstruction failed")
        
        structure = DocumentStructure(format="docx")
        dims = Dimensions(width=612, height=792)
        page = PageStructure(page_number=1, dimensions=dims)
        structure.add_page(page)
        
        with pytest.raises(ReconstructionError, match="Failed to reconstruct DOCX"):
            parser.reconstruct(structure)
    
    def test_apply_formatting_to_run(self):
        """Test applying formatting to a run."""
        parser = DOCXParser()
        
        # Mock run and font
        mock_run = Mock()
        mock_font = Mock()
        mock_run.font = mock_font
        
        # Create formatting
        formatting = TextFormatting(
            font_family="Arial",
            font_size=14,
            is_bold=True,
            is_italic=True,
            is_underlined=True,
            color="#FF0000"
        )
        
        # Mock Pt and RGBColor
        with patch('src.parsers.docx_parser.Pt') as mock_pt, \
             patch('docx.shared.RGBColor') as mock_rgb_color:
            
            mock_pt.return_value = "14pt"
            mock_rgb_color.return_value = "red_color"
            
            parser._apply_formatting_to_run(mock_run, formatting)
        
        # Verify formatting was applied
        assert mock_font.name == "Arial"
        mock_pt.assert_called_with(14)
        assert mock_font.bold is True
        assert mock_font.italic is True
        assert mock_font.underline is True
        mock_rgb_color.assert_called_with(255, 0, 0)
    
    def test_document_to_bytes(self):
        """Test converting document to bytes."""
        parser = DOCXParser()
        
        # Mock document
        mock_doc = Mock()
        
        # Mock BytesIO
        mock_bytes_io = Mock()
        mock_bytes_io.read.return_value = b"docx file content"
        
        with patch('io.BytesIO', return_value=mock_bytes_io):
            result = parser._document_to_bytes(mock_doc)
        
        assert result == b"docx file content"
        mock_doc.save.assert_called_once_with(mock_bytes_io)
        mock_bytes_io.seek.assert_called_once_with(0)
        mock_bytes_io.read.assert_called_once()
    
    def test_extract_visual_elements(self):
        """Test visual element extraction."""
        parser = DOCXParser()
        
        # Mock document with image relationships
        mock_doc = Mock()
        
        # Mock relationship with image
        mock_rel = Mock()
        mock_rel.target_ref = "image1.png"
        mock_rel.target_part.blob = b"fake image data"
        mock_rel.target_part.content_type = "image/png"
        
        mock_doc.part.rels.values.return_value = [mock_rel]
        
        visual_elements = parser._extract_visual_elements(mock_doc)
        
        assert len(visual_elements) == 1
        element = visual_elements[0]
        assert element.element_type == "image"
        assert element.content == b"fake image data"
        assert element.metadata["content_type"] == "image/png"
        assert element.metadata["filename"] == "image1.png"
    
    def test_build_spatial_map(self):
        """Test spatial map building."""
        parser = DOCXParser()
        
        # Create test elements
        text_region1 = TextRegion(
            id="text1",
            bounding_box=BoundingBox(x=10, y=10, width=50, height=20)
        )
        text_region2 = TextRegion(
            id="text2",
            bounding_box=BoundingBox(x=10, y=40, width=50, height=20)
        )
        visual_element = VisualElement(
            id="image1",
            element_type="image",
            bounding_box=BoundingBox(x=70, y=10, width=30, height=30)
        )
        
        spatial_map = parser._build_spatial_map(
            [text_region1, text_region2],
            [visual_element]
        )
        
        # Check reading order (sorted by y, then x)
        assert spatial_map.reading_order == ["text1", "image1", "text2"]
        
        # Check relationships (elements within distance threshold)
        assert "text1" in spatial_map.element_relationships
        assert "text2" in spatial_map.element_relationships["text1"]