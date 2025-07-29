"""DOCX document parser using python-docx."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import logging
from datetime import datetime
import io

from .base import DocumentParser, ParsingError, ReconstructionError
from src.models.document import (
    DocumentStructure,
    PageStructure,
    TextRegion,
    VisualElement,
    BoundingBox,
    Dimensions,
    TextFormatting,
    DocumentMetadata,
    SpatialMap,
)


class DOCXParser(DocumentParser):
    """DOCX document parser using python-docx for Word document processing."""

    def __init__(self):
        """Initialize the DOCX parser."""
        super().__init__()
        self.supported_formats = ["docx"]
        # Standard page dimensions for Letter size (8.5" x 11")
        self.default_page_width = 612  # points
        self.default_page_height = 792  # points

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats.

        Returns:
            List containing 'docx'
        """
        return self.supported_formats

    def parse(self, file_path: str) -> DocumentStructure:
        """Parse a DOCX document and extract its structure.

        Args:
            file_path: Path to the DOCX file

        Returns:
            DocumentStructure containing parsed content and layout

        Raises:
            ParsingError: If DOCX parsing fails
        """
        try:
            self.logger.info(f"Starting DOCX parsing: {file_path}")

            # Open the DOCX document
            doc = Document(file_path)

            # Extract document metadata
            metadata = self._extract_docx_metadata(doc, file_path)

            # Create document structure
            document = DocumentStructure(format="docx", metadata=metadata)

            # Process document content
            # Note: DOCX doesn't have explicit pages like PDF, so we simulate pages
            page_structure = self._parse_document_content(doc)
            document.add_page(page_structure)

            self.logger.info(
                f"Successfully parsed DOCX: {file_path} "
                f"({len(page_structure.text_regions)} text regions, "
                f"{len(page_structure.visual_elements)} visual elements)"
            )

            return document

        except Exception as e:
            if isinstance(e, ParsingError):
                raise
            raise ParsingError(
                f"Failed to parse DOCX document: {str(e)}",
                file_path,
                "DOCX_PARSE_ERROR",
            )

    def reconstruct(self, structure: DocumentStructure) -> bytes:
        """Reconstruct a DOCX document from its structure.

        Args:
            structure: DocumentStructure to reconstruct

        Returns:
            Binary content of the reconstructed DOCX

        Raises:
            ReconstructionError: If DOCX reconstruction fails
        """
        try:
            self.logger.info(
                f"Starting DOCX reconstruction ({len(structure.pages)} pages)"
            )

            # Create new DOCX document
            doc = Document()

            # Set document metadata
            self._set_document_metadata(doc, structure.metadata)

            # Reconstruct content from all pages
            for page_structure in structure.pages:
                self._reconstruct_page_content(doc, page_structure)

            # Save to bytes
            docx_bytes = self._document_to_bytes(doc)

            self.logger.info(
                f"Successfully reconstructed DOCX ({len(docx_bytes)} bytes)"
            )

            return docx_bytes

        except Exception as e:
            raise ReconstructionError(
                f"Failed to reconstruct DOCX: {str(e)}",
                "docx",
                "DOCX_RECONSTRUCTION_ERROR",
            )

    def _extract_docx_metadata(self, doc: Document, file_path: str) -> DocumentMetadata:
        """Extract metadata from DOCX document.

        Args:
            doc: python-docx Document object
            file_path: Path to the DOCX file

        Returns:
            DocumentMetadata with extracted information
        """
        core_props = doc.core_properties
        file_stat = Path(file_path).stat()

        return DocumentMetadata(
            title=core_props.title or Path(file_path).stem,
            author=core_props.author,
            subject=core_props.subject,
            creator=core_props.creator,
            creation_date=core_props.created,
            modification_date=core_props.modified,
            page_count=1,  # Will be updated based on content
            file_size=file_stat.st_size,
        )

    def _parse_document_content(self, doc: Document) -> PageStructure:
        """Parse the content of a DOCX document.

        Args:
            doc: python-docx Document object

        Returns:
            PageStructure with extracted content
        """
        # Create page structure (DOCX is treated as single logical page)
        dimensions = Dimensions(
            width=self.default_page_width, height=self.default_page_height, unit="pt"
        )

        page_structure = PageStructure(page_number=1, dimensions=dimensions)

        # Track vertical position for layout
        current_y = 50  # Start with top margin

        # Process paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Skip empty paragraphs
                text_regions = self._process_paragraph(paragraph, para_idx, current_y)
                page_structure.text_regions.extend(text_regions)

                # Update vertical position (approximate)
                current_y += len(text_regions) * 20  # Approximate line height

        # Process tables
        for table_idx, table in enumerate(doc.tables):
            table_elements = self._process_table(table, table_idx, current_y)
            page_structure.text_regions.extend(table_elements)
            current_y += len(table.rows) * 25  # Approximate row height

        # Process images and other elements
        visual_elements = self._extract_visual_elements(doc)
        page_structure.visual_elements.extend(visual_elements)

        # Build spatial map
        spatial_map = self._build_spatial_map(
            page_structure.text_regions, page_structure.visual_elements
        )
        page_structure.spatial_map = spatial_map

        return page_structure

    def _process_paragraph(
        self, paragraph, para_idx: int, y_position: float
    ) -> List[TextRegion]:
        """Process a paragraph and extract text regions.

        Args:
            paragraph: python-docx Paragraph object
            para_idx: Paragraph index
            y_position: Vertical position on page

        Returns:
            List of TextRegion objects
        """
        text_regions = []
        current_x = 50  # Left margin

        # Process runs (text with consistent formatting)
        for run_idx, run in enumerate(paragraph.runs):
            if not run.text.strip():
                continue

            # Estimate text width (approximate)
            font_size = run.font.size.pt if run.font.size else 12
            text_width = len(run.text) * font_size * 0.6  # Rough estimation

            # Create bounding box
            bbox = BoundingBox(
                x=current_x,
                y=y_position,
                width=text_width,
                height=font_size * 1.2,  # Line height
            )

            # Extract formatting
            formatting = self._extract_run_formatting(run, paragraph)

            # Create text region
            region_id = f"para_{para_idx}_run_{run_idx}"
            text_region = TextRegion(
                id=region_id,
                bounding_box=bbox,
                text_content=run.text,
                formatting=formatting,
                language="en",  # Will be detected later
                confidence=1.0,
                reading_order=len(text_regions),
            )

            text_regions.append(text_region)
            current_x += text_width

        return text_regions

    def _extract_run_formatting(self, run, paragraph) -> TextFormatting:
        """Extract formatting from a run and paragraph.

        Args:
            run: python-docx Run object
            paragraph: python-docx Paragraph object

        Returns:
            TextFormatting object
        """
        # Get font properties
        font = run.font
        font_size = font.size.pt if font.size else 12
        font_name = font.name or "Calibri"

        # Get text properties
        is_bold = font.bold if font.bold is not None else False
        is_italic = font.italic if font.italic is not None else False
        is_underlined = font.underline if font.underline is not None else False

        # Get color (if available)
        color = "#000000"  # Default black
        if font.color and font.color.rgb:
            rgb = font.color.rgb
            color = f"#{rgb.r:02X}{rgb.g:02X}{rgb.b:02X}"

        # Get alignment
        alignment = "left"  # Default
        if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            alignment = "center"
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            alignment = "right"
        elif paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            alignment = "justify"

        return TextFormatting(
            font_family=font_name,
            font_size=font_size,
            is_bold=is_bold,
            is_italic=is_italic,
            is_underlined=is_underlined,
            color=color,
            alignment=alignment,
        )

    def _process_table(
        self, table, table_idx: int, y_position: float
    ) -> List[TextRegion]:
        """Process a table and extract text regions from cells.

        Args:
            table: python-docx Table object
            table_idx: Table index
            y_position: Vertical position on page

        Returns:
            List of TextRegion objects from table cells
        """
        text_regions = []
        cell_height = 25  # Approximate cell height
        cell_width = 100  # Approximate cell width

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if cell.text.strip():
                    # Calculate cell position
                    x_pos = 50 + col_idx * cell_width
                    y_pos = y_position + row_idx * cell_height

                    bbox = BoundingBox(
                        x=x_pos, y=y_pos, width=cell_width, height=cell_height
                    )

                    # Use default formatting for table cells
                    formatting = TextFormatting(
                        font_family="Calibri", font_size=11, alignment="left"
                    )

                    region_id = f"table_{table_idx}_row_{row_idx}_col_{col_idx}"
                    text_region = TextRegion(
                        id=region_id,
                        bounding_box=bbox,
                        text_content=cell.text,
                        formatting=formatting,
                        language="en",
                        confidence=1.0,
                        reading_order=len(text_regions),
                    )

                    text_regions.append(text_region)

        return text_regions

    def _extract_visual_elements(self, doc: Document) -> List[VisualElement]:
        """Extract visual elements from DOCX document.

        Args:
            doc: python-docx Document object

        Returns:
            List of VisualElement objects
        """
        visual_elements = []

        # Extract inline shapes (images, charts, etc.)
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    # Get image data
                    image_data = rel.target_part.blob

                    # Create visual element (position is approximate)
                    element_id = f"image_{len(visual_elements)}"
                    bbox = BoundingBox(x=100, y=100, width=200, height=150)

                    visual_element = VisualElement(
                        id=element_id,
                        element_type="image",
                        bounding_box=bbox,
                        content=image_data,
                        metadata={
                            "content_type": rel.target_part.content_type,
                            "filename": rel.target_ref,
                        },
                    )

                    visual_elements.append(visual_element)

                except Exception as e:
                    self.logger.warning(f"Failed to extract image: {str(e)}")

        return visual_elements

    def _build_spatial_map(
        self, text_regions: List[TextRegion], visual_elements: List[VisualElement]
    ) -> SpatialMap:
        """Build spatial relationships map for document elements.

        Args:
            text_regions: List of text regions
            visual_elements: List of visual elements

        Returns:
            SpatialMap with element relationships
        """
        spatial_map = SpatialMap()

        # Set reading order based on document flow (top to bottom)
        all_elements = text_regions + visual_elements
        sorted_elements = sorted(
            all_elements, key=lambda elem: (elem.bounding_box.y, elem.bounding_box.x)
        )

        spatial_map.reading_order = [elem.id for elem in sorted_elements]

        # Build relationships based on proximity
        for i, element in enumerate(all_elements):
            nearby_elements = []

            for j, other_element in enumerate(all_elements):
                if i == j:
                    continue

                # Check if elements are nearby
                distance = self._calculate_distance(
                    element.bounding_box, other_element.bounding_box
                )

                if distance < 150:  # Threshold for "nearby" in DOCX
                    nearby_elements.append(other_element.id)

            if nearby_elements:
                spatial_map.add_relationship(element.id, nearby_elements)

        return spatial_map

    def _calculate_distance(self, bbox1: BoundingBox, bbox2: BoundingBox) -> float:
        """Calculate distance between two bounding boxes.

        Args:
            bbox1: First bounding box
            bbox2: Second bounding box

        Returns:
            Distance between box centers
        """
        center1_x = bbox1.x + bbox1.width / 2
        center1_y = bbox1.y + bbox1.height / 2
        center2_x = bbox2.x + bbox2.width / 2
        center2_y = bbox2.y + bbox2.height / 2

        return ((center1_x - center2_x) ** 2 + (center1_y - center2_y) ** 2) ** 0.5

    def _set_document_metadata(self, doc: Document, metadata: DocumentMetadata) -> None:
        """Set document metadata.

        Args:
            doc: python-docx Document object
            metadata: DocumentMetadata to set
        """
        core_props = doc.core_properties

        if metadata.title:
            core_props.title = metadata.title
        if metadata.author:
            core_props.author = metadata.author
        if metadata.subject:
            core_props.subject = metadata.subject
        if metadata.creator:
            core_props.creator = metadata.creator

    def _reconstruct_page_content(
        self, doc: Document, page_structure: PageStructure
    ) -> None:
        """Reconstruct page content in DOCX document.

        Args:
            doc: python-docx Document object
            page_structure: Page structure to reconstruct
        """
        # Group text regions by their vertical position to form paragraphs
        text_groups = self._group_text_regions_by_position(page_structure.text_regions)

        # Add paragraphs to document
        for group in text_groups:
            paragraph = doc.add_paragraph()

            for text_region in group:
                run = paragraph.add_run(text_region.text_content)
                self._apply_formatting_to_run(run, text_region.formatting)

            # Set paragraph alignment
            if group and group[0].formatting.alignment == "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif group and group[0].formatting.alignment == "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif group and group[0].formatting.alignment == "justify":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Add visual elements (simplified - just add a placeholder)
        for visual_element in page_structure.visual_elements:
            if visual_element.element_type == "image" and visual_element.content:
                try:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()

                    # Add image from bytes
                    image_stream = io.BytesIO(visual_element.content)
                    run.add_picture(image_stream, width=Inches(3))

                except Exception as e:
                    self.logger.warning(
                        f"Failed to add image {visual_element.id}: {str(e)}"
                    )

    def _group_text_regions_by_position(
        self, text_regions: List[TextRegion]
    ) -> List[List[TextRegion]]:
        """Group text regions by their vertical position to form paragraphs.

        Args:
            text_regions: List of text regions

        Returns:
            List of text region groups (paragraphs)
        """
        if not text_regions:
            return []

        # Sort by vertical position
        sorted_regions = sorted(text_regions, key=lambda r: r.bounding_box.y)

        groups = []
        current_group = [sorted_regions[0]]
        current_y = sorted_regions[0].bounding_box.y

        for region in sorted_regions[1:]:
            # If regions are close vertically, they're likely in the same paragraph
            if abs(region.bounding_box.y - current_y) < 25:  # Same line threshold
                current_group.append(region)
            else:
                # Start new paragraph
                groups.append(current_group)
                current_group = [region]
                current_y = region.bounding_box.y

        # Add the last group
        if current_group:
            groups.append(current_group)

        return groups

    def _apply_formatting_to_run(self, run, formatting: TextFormatting) -> None:
        """Apply formatting to a run.

        Args:
            run: python-docx Run object
            formatting: TextFormatting to apply
        """
        font = run.font

        if formatting.font_family:
            font.name = formatting.font_family

        if formatting.font_size:
            font.size = Pt(formatting.font_size)

        if formatting.is_bold:
            font.bold = True

        if formatting.is_italic:
            font.italic = True

        if formatting.is_underlined:
            font.underline = True

        # Set color if not default black
        if formatting.color and formatting.color != "#000000":
            try:
                # Convert hex to RGB
                hex_color = formatting.color.lstrip("#")
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)

                from docx.shared import RGBColor

                font.color.rgb = RGBColor(r, g, b)
            except (ValueError, IndexError):
                pass  # Use default color

    def _document_to_bytes(self, doc: Document) -> bytes:
        """Convert Document object to bytes.

        Args:
            doc: python-docx Document object

        Returns:
            Document content as bytes
        """
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.read()
